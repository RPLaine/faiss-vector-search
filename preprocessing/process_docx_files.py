"""
DOCX File Processor with Markdown Conversion

This script processes all DOCX files in the 'unprocessed' directory,
extracts their content while preserving formatting, and converts
it to Markdown format before saving as text files in the 'txt-files' directory.

Usage:
    python process_docx_files.py

The script will:
1. Scan the 'unprocessed' directory for all .docx files
2. Extract text content and formatting from each DOCX file
3. Convert formatting to Markdown (headers, bold, italic, lists, etc.)
4. Save as .txt files in the 'txt-files' directory with Markdown formatting
5. Preserve document structure and handle encoding gracefully

Markdown conversions:
- Headings → # ## ### #### #####
- Bold text → **bold**
- Italic text → *italic*
- Lists → - item or 1. item
- Tables → Markdown table format
"""

from pathlib import Path
from docx import Document
import re
import shutil


def clear_directory(directory_path):
    """Clear all files from a directory while keeping the directory itself."""
    if directory_path.exists():
        for item in directory_path.iterdir():
            if item.is_file():
                item.unlink()  # Remove file
            elif item.is_dir():
                shutil.rmtree(item)  # Remove directory and contents
        print(f"  Cleared existing files from {directory_path}")


def get_heading_level(paragraph):
    """Determine the heading level of a paragraph based on its style."""
    style_name = paragraph.style.name.lower()
    if 'heading' in style_name:
        # Extract number from style name like "Heading 1", "Heading 2", etc.
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
    return 0


def format_run_text(run):
    """Convert a run's formatting to markdown."""
    text = run.text
    if not text.strip():
        return text
    
    # Apply bold formatting
    if run.bold:
        text = f"**{text}**"
    
    # Apply italic formatting  
    if run.italic:
        text = f"*{text}*"
    
    return text


def process_paragraph(paragraph):
    """Convert a paragraph to markdown format."""
    if not paragraph.text.strip():
        return ""
    
    # Check if it's a heading
    heading_level = get_heading_level(paragraph)
    if heading_level > 0:
        heading_prefix = "#" * min(heading_level, 6)  # Limit to 6 levels
        return f"{heading_prefix} {paragraph.text.strip()}\n"
    
    # Check if it's a list item
    if paragraph.style.name.startswith('List'):
        # Simple bullet point for now
        return f"- {paragraph.text.strip()}\n"
    
    # Process runs for inline formatting
    formatted_text = ""
    for run in paragraph.runs:
        formatted_text += format_run_text(run)
    
    return formatted_text.strip() + "\n"


def process_table(table):
    """Convert a table to markdown format."""
    if not table.rows:
        return ""
    
    markdown_table = []
    
    # Process header row
    header_row = table.rows[0]
    header_cells = [cell.text.strip() for cell in header_row.cells]
    markdown_table.append("| " + " | ".join(header_cells) + " |")
    
    # Add separator row
    separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"
    markdown_table.append(separator)
    
    # Process data rows
    for row in table.rows[1:]:
        row_cells = [cell.text.strip() for cell in row.cells]
        # Pad row with empty cells if needed
        while len(row_cells) < len(header_cells):
            row_cells.append("")
        markdown_table.append("| " + " | ".join(row_cells[:len(header_cells)]) + " |")
    
    return "\n".join(markdown_table) + "\n\n"


def convert_docx_to_markdown(docx_path):
    """Convert a DOCX file to markdown format."""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        md_text = process_paragraph(paragraph)
                        if md_text:
                            markdown_content.append(md_text)
                        break
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        md_table = process_table(table)
                        if md_table:
                            markdown_content.append(md_table)
                        break
        
        # Join content and clean up excessive newlines
        result = "".join(markdown_content)
        # Replace multiple consecutive newlines with double newlines
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result.strip()
        
    except Exception as e:
        print(f"Error converting {docx_path} to markdown: {str(e)}")
        return ""


def process_docx_files():
    """
    Process all DOCX files in the unprocessed directory and save them as markdown text files
    in the txt-files directory.
    """
    # Define directories relative to project root
    script_dir = Path(__file__).parent.parent  # Go up to project root
    docx_dir = script_dir / "unprocessed"
    output_dir = script_dir / "txt-files"
    
    # Create output directory if it doesn't exist
    output_dir.mkdir(exist_ok=True)
    
    # Clear existing files from output directory
    clear_directory(output_dir)
    
    # Check if unprocessed directory exists
    if not docx_dir.exists():
        print(f"Error: Directory '{docx_dir}' does not exist!")
        return
    
    # Get all DOCX files
    docx_files = list(docx_dir.glob("*.docx"))
    
    if not docx_files:
        print(f"No DOCX files found in '{docx_dir}' directory!")
        return
    
    print(f"Found {len(docx_files)} DOCX files to process...")
    
    successful_count = 0
    failed_count = 0
    
    for docx_file in docx_files:
        try:
            print(f"Processing: {docx_file.name}")
            
            # Convert DOCX to markdown
            markdown_content = convert_docx_to_markdown(str(docx_file))
            
            if not markdown_content.strip():
                print(f"  ⚠ No content extracted from {docx_file.name}")
                continue
            
            # Create output filename (replace .docx with .txt)
            output_filename = docx_file.stem + ".txt"
            output_path = output_dir / output_filename
            
            # Save markdown content to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"  ✓ Saved to: {output_path}")
            successful_count += 1
            
        except Exception as e:
            print(f"  ✗ Error processing {docx_file.name}: {str(e)}")
            failed_count += 1
    
    # Print summary
    print(f"\nProcessing complete!")
    print(f"Successfully processed: {successful_count} files")
    if failed_count > 0:
        print(f"Failed to process: {failed_count} files")
    
    print(f"Markdown text files saved in: {output_dir.absolute()}")


if __name__ == "__main__":
    process_docx_files()